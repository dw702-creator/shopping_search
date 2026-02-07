import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# ---------------- Word 시험지 생성 ----------------
def create_exam_doc(text):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    # 제목
    title = doc.add_paragraph("시험지")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(16)

    info = doc.add_paragraph("반: ________   이름: ________   점수: ________   선생님 확인: ________")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.runs[0].font.size = Pt(11)

    doc.add_paragraph("")

    # 2열 테이블 (문제 / 메모)
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(3.5)
    table.columns[1].width = Inches(3.5)

    # 입력 텍스트를 문단 단위로 분리
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    for i, para in enumerate(paragraphs, start=1):
        p = left_cell.add_paragraph(f"{i}. {para}")
        p.runs[0].font.size = Pt(11)

        memo = right_cell.add_paragraph("\n\n\n")  # 메모 공간
        memo.runs[0].font.size = Pt(11)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------------- Streamlit UI ----------------
st.set_page_config(page_title="시험지 생성기", layout="wide")
st.title("📝 시험지 생성기 (문제 + 메모 공간)")

input_text = st.text_area(
    "시험지로 만들 텍스트를 입력하세요",
    height=300,
    placeholder="여기에 문제로 사용할 텍스트를 입력하세요.\n문단 단위로 나뉩니다."
)

if input_text.strip():
    if st.button("시험지 Word 파일 생성"):
        file = create_exam_doc(input_text)
        st.download_button(
            label="⬇️ 시험지 다운로드 (.docx)",
            data=file,
            file_name="시험지_문제+메모.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
else:
    st.info("텍스트를 입력하면 시험지를 생성할 수 있습니다.")
